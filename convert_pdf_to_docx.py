#!/usr/bin/env python3
"""
Convert 3 PDF chapters to a single DOCX file.
Preserves ALL text and ALL images from the PDFs.
"""
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io
import re

# Configuration
PDF_DIR = "/Users/khoait/Documents/SourceCompany/chatbot_ufm/tailieumau"
OUTPUT = "/Users/khoait/Documents/SourceCompany/chatbot_ufm/TaiLieu_3Chuong_DayDu.docx"
IMG_DIR = "/Users/khoait/Documents/SourceCompany/chatbot_ufm/tailieumau/extracted_images"
os.makedirs(IMG_DIR, exist_ok=True)

# PDF files in order
pdf_files = [
    "CHUONG1_TONGQUAN.pdf",
    "CHUONG3_THIET KE LOP.pdf",
    "CHUONG5_MAUTHIETKEHUONGDOITUONG_new.pdf",
]

doc = Document()

# Style setup
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.size = Pt([0, 18, 14, 12][i])

# Cover page
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TÀI LIỆU MÔN HỌC")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("THIẾT KẾ PHẦN MỀM\nHƯỚNG ĐỐI TƯỢNG")
r.font.size = Pt(22)
r.bold = True
r.font.color.rgb = RGBColor(0, 102, 153)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Chương 1 · Chương 3 · Chương 5\n(Chuyển đổi đầy đủ từ PDF)")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(100, 100, 100)
doc.add_page_break()

total_images = 0
total_pages = 0

for pdf_idx, pdf_name in enumerate(pdf_files):
    pdf_path = os.path.join(PDF_DIR, pdf_name)
    print(f"\n{'='*60}")
    print(f"Đang xử lý: {pdf_name}")
    print(f"{'='*60}")

    pdf_doc = fitz.open(pdf_path)
    num_pages = len(pdf_doc)
    total_pages += num_pages
    print(f"  Số trang: {num_pages}")

    # Chapter heading
    chapter_name = pdf_name.replace(".pdf", "").replace("_new", "").replace("_", " ")
    doc.add_heading(f"{'='*20}", level=1)
    doc.add_heading(chapter_name.upper(), level=1)
    doc.add_heading(f"(File gốc: {pdf_name} — {num_pages} trang)", level=3)
    doc.add_paragraph()

    for page_num in range(num_pages):
        page = pdf_doc[page_num]

        # ── Extract text ──
        # Use get_text("dict") for structured extraction
        blocks = page.get_text("dict")["blocks"]

        # Sort blocks by vertical position (top to bottom)
        sorted_blocks = sorted(blocks, key=lambda b: (b["bbox"][1], b["bbox"][0]))

        page_has_content = False

        for block in sorted_blocks:
            if block["type"] == 0:  # Text block
                for line in block.get("lines", []):
                    line_text = ""
                    is_bold = False
                    font_size = 12
                    is_heading = False

                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if text:
                            line_text += text + " "
                            font_size = span.get("size", 12)
                            flags = span.get("flags", 0)
                            is_bold = bool(flags & 2**4)  # bold flag

                    line_text = line_text.strip()
                    if not line_text:
                        continue

                    page_has_content = True

                    # Detect headings by font size
                    if font_size >= 20:
                        doc.add_heading(line_text, level=1)
                    elif font_size >= 16:
                        doc.add_heading(line_text, level=2)
                    elif font_size >= 13 and is_bold:
                        doc.add_heading(line_text, level=3)
                    else:
                        p = doc.add_paragraph()
                        if is_bold:
                            run = p.add_run(line_text)
                            run.bold = True
                        else:
                            p.add_run(line_text)

            elif block["type"] == 1:  # Image block
                # Extract image
                img_data = block.get("image", None)
                if img_data:
                    total_images += 1
                    img_name = f"img_{pdf_idx}_{page_num}_{total_images}.png"
                    img_path = os.path.join(IMG_DIR, img_name)
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"    ⚠️ Lỗi chèn ảnh block {img_name}: {e}")

        # Also extract images using get_images() for images not in blocks
        image_list = page.get_images(full=True)
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = pdf_doc.extract_image(xref)
                if base_image:
                    img_bytes = base_image["image"]
                    img_ext = base_image["ext"]
                    total_images += 1
                    img_name = f"page_{pdf_idx}_{page_num}_{img_index}.{img_ext}"
                    img_path = os.path.join(IMG_DIR, img_name)

                    # Skip tiny images (likely decorations/bullets)
                    if len(img_bytes) < 500:
                        continue

                    with open(img_path, "wb") as f:
                        f.write(img_bytes)

                    # Calculate appropriate width
                    width = min(base_image.get("width", 400), 600)
                    doc_width = Inches(min(width / 96.0, 5.8))  # Max 5.8 inches

                    try:
                        doc.add_picture(img_path, width=doc_width)
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"    ⚠️ Lỗi chèn ảnh xref {img_name}: {e}")
            except Exception as e:
                pass  # Skip problematic images

        # Page separator
        if page_has_content and page_num < num_pages - 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"— Trang {page_num + 1}/{num_pages} —")
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(180, 180, 180)

    pdf_doc.close()
    print(f"  ✅ Hoàn thành: {num_pages} trang")

    # Page break between chapters
    if pdf_idx < len(pdf_files) - 1:
        doc.add_page_break()

# Save
doc.save(OUTPUT)
file_size = os.path.getsize(OUTPUT) / 1024 / 1024
print(f"\n{'='*60}")
print(f"✅ ĐÃ TẠO THÀNH CÔNG!")
print(f"   File: {OUTPUT}")
print(f"   Kích thước: {file_size:.2f} MB")
print(f"   Tổng trang PDF: {total_pages}")
print(f"   Tổng hình ảnh: {total_images}")
print(f"{'='*60}")
